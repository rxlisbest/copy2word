#!/usr/bin/python
# -*- coding: utf-8 -*-

import os
from docx import Document
from docx.shared import Inches
from docx.shared import Pt

document = Document('./模板.docx')
name = '养老大数据平台_v1.5.0'
document.sections[0].header.paragraphs[0].text = name
# document.sections[0].footer.paragraphs[0].text= '党建规范化可视应用平台运营版V1.0'

model_dir = '/Library/WebServer/Documents/bigdata/apps/'

def getFile(dir, files):
  l = os.listdir(dir)
  for i in l:
    if os.path.isdir(dir + i): 
      files = getFile(dir + i + '/', files)
    elif os.path.isfile(dir + i):
      files.append(dir + i)
    else:
      print(i)
  return files
files = getFile(model_dir, []) 

num = 0
for i in files:
  fp = open(i)
  for line in fp.readlines(): 
    if line.strip().replace('\n', '').replace('\r', '').replace('\r\n', '') != '':
      # print(line)
      paragraph = document.add_paragraph()
      run = paragraph.add_run(line.replace('\n', '').replace('\r', '').replace('\r\n', ''))
      run.font.name = 'Menlo'
      run.font.size = Pt(9)
      paragraph_format = paragraph.paragraph_format
      paragraph_format.space_before = Pt(0)
      paragraph_format.space_after = Pt(0)
      num += 1
  if num > 4000:
    break
  fp.close()

document.save("./%s.docx" % name)
